"""Generates the assignment write-up as a .docx file."""
import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

HERE = os.path.dirname(os.path.abspath(__file__))
FIG = os.path.join(HERE, "figures")
OUT = os.path.join(HERE, "JVM-Tuning-Assignment.docx")

doc = Document()

# The default template uses 1.25 in side margins, which leaves only 6.00 in of
# printable width. Tighten to 1 in so the wide screenshots fit comfortably.
for section in doc.sections:
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

PRINTABLE_IN = 6.5

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(8)
style.paragraph_format.line_spacing = 1.15

_fig_no = [0]


def shade(paragraph, hex_fill):
    pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    pr.append(shd)


def code_block(lines):
    for i, line in enumerate(lines):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0) if i < len(lines) - 1 else Pt(10)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.left_indent = Inches(0.25)
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(9.5)
        shade(p, "F2F2F2")


def figure(filename, caption, width=6.3):
    """Insert a real screenshot with a numbered caption."""
    assert width <= PRINTABLE_IN, f"{filename} at {width}in exceeds printable width"
    _fig_no[0] += 1
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(os.path.join(FIG, filename), width=Inches(width))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(14)
    run = cap.add_run(f"Figure {_fig_no[0]}: {caption}")
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        shade(hdr[i].paragraphs[0], "DDDDDD")
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


# ----------------------------------------------------------------- title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
trun = title.add_run("JVM Tuning and Spring Boot Microservice Performance")
trun.bold = True
trun.font.size = Pt(18)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after = Pt(18)
for text, br in [("Name: [Your Name]", True),
                 ("Course / Section: [Course]", True),
                 ("Date: [Date]", False)]:
    r = sub.add_run(text)
    r.font.size = Pt(10.5)
    if br:
        r.add_break()

# ----------------------------------------------------------- introduction
doc.add_heading("1. Introduction", level=1)
doc.add_paragraph(
    "A microservice is a small application that does one job and communicates over HTTP. "
    "I used Spring Boot to build one because it comes with an embedded Tomcat server and "
    "sets up most of the configuration on its own, so a working REST endpoint only takes a "
    "few lines of code."
)
doc.add_paragraph(
    "Every Spring Boot application runs inside the JVM, and the JVM is what decides how much "
    "memory the application can use and when it throws away objects that are no longer needed. "
    "That cleanup is called garbage collection. When the heap is too small for the workload, the "
    "collector has to run over and over, and the application ends up spending more time freeing "
    "memory than answering requests. JVM tuning is the process of setting options such as the "
    "heap size and the collector type so the application matches the work it actually does."
)
doc.add_paragraph(
    "For this assignment I built a small item service, put it under load, watched it in VisualVM, "
    "then changed the JVM settings and measured the difference."
)
p = doc.add_paragraph()
r = p.add_run("Environment: ")
r.bold = True
p.add_run("Java 21.0.2 (OpenJDK), Spring Boot 3.4.1, Maven 3.9, VisualVM 2.2, "
          "8 CPU cores, 31 GB RAM.")

# --------------------------------------------------------------- step 1
doc.add_heading("2. Step 1: Building the Microservice", level=1)
doc.add_paragraph(
    "I created a Maven project with the Spring Web and Spring Boot Actuator dependencies. The "
    "whole service is three classes. The controller is the only part that really matters:"
)
code_block([
    "@RestController",
    "@RequestMapping(\"/api/items\")",
    "public class ItemController {",
    "",
    "    private static final String[] CATEGORIES =",
    "            {\"Tools\", \"Books\", \"Food\", \"Toys\", \"Parts\"};",
    "",
    "    @GetMapping",
    "    public List<Item> getItems(@RequestParam(defaultValue = \"5000\") int count) {",
    "        List<Item> items = new ArrayList<>();",
    "        for (int i = 0; i < count; i++) {",
    "            String name = \"Item \" + i + \" - \" + CATEGORIES[i % CATEGORIES.length];",
    "            items.add(new Item(i, name,",
    "                    CATEGORIES[i % CATEGORIES.length], 10.0 + (i % 100)));",
    "        }",
    "        return items;",
    "    }",
    "}",
])
doc.add_paragraph(
    "Item itself is just a record with an id, name, category and price. The endpoint is "
    "GET /api/items?count=5000 and it returns the list as JSON."
)
doc.add_paragraph(
    "One thing I did on purpose: the list is rebuilt on every single request instead of being "
    "cached. If I cached it the service would barely allocate anything and there would be nothing "
    "interesting to look at in VisualVM. Rebuilding it means every request creates thousands of "
    "short lived objects, which is what puts pressure on the garbage collector."
)
doc.add_paragraph("Building and running it:")
code_block([
    "mvn package -DskipTests",
    "java -jar target/itemservice-0.0.1-SNAPSHOT.jar",
])
doc.add_paragraph(
    "I checked the endpoint worked before worrying about performance. A full 5,000 item response "
    "comes back as 349 KB of JSON in about 43 ms when the service is not under load."
)
figure("fig1-terminal.png",
       "The endpoint returning JSON, the size and timing of a 5,000 item response, "
       "and the Actuator health check", width=5.9)

# --------------------------------------------------------------- step 2
doc.add_heading("3. Step 2: Monitoring with VisualVM", level=1)
doc.add_paragraph(
    "I downloaded VisualVM from visualvm.github.io and started it while the service was already "
    "running. The local Java process shows up in the panel on the left, listed by its main class. "
    "Double clicking it opens the application, and the Overview tab is a good place to start "
    "because it confirms which JVM options the process is actually running with."
)
doc.add_paragraph(
    "For the first run I started the service with a deliberately small heap and the serial "
    "collector, so the problem would be easy to see:"
)
code_block(["java -Xms64m -Xmx128m -XX:+UseSerialGC \\",
            "     -cp \"target/classes:$(cat cp.txt)\" \\",
            "     com.example.itemservice.ItemserviceApplication"])
figure("fig-baseline-overview.png",
       "Overview tab for the baseline run. The JVM arguments panel confirms "
       "-Xms64m, -Xmx128m and -XX:+UseSerialGC")

doc.add_paragraph(
    "An idle application does not show much, so I needed traffic. I wrote a short Python script "
    "that sends requests with 16 threads asking for 25,000 items each, and I ran it in a loop so "
    "the load stayed constant for about two and a half minutes while I watched the Monitor tab."
)
figure("fig-baseline-monitor.png",
       "Monitor tab during the baseline run. Heap is pinned at its 128 MB ceiling and used heap "
       "sawtooths violently between roughly 25 MB and 100 MB")

doc.add_paragraph("What the Monitor tab showed once the load started:")
for line in [
    "Heap size jumped straight to the ceiling and stayed there. Size and Max both read "
    "134,217,728 B, which is the 128 MB I asked for.",
    "Used heap sawtoothed between roughly 25 MB and 100 MB and never settled, so the collector "
    "was reclaiming memory continuously rather than in occasional bursts.",
    "CPU usage sat around 34 percent and GC activity was reported at 7.2 percent.",
    "Live threads went from 29 to 37 when the load arrived and then stayed flat, so thread count "
    "was not the problem.",
]:
    doc.add_paragraph(line, style="List Bullet")

doc.add_heading("Why the GC activity percentage looked deceptively small", level=2)
doc.add_paragraph(
    "Reading exact pause times off a moving graph is guesswork, so I also turned on GC logging:"
)
code_block(["-Xlog:gc:file=baseline-gc.log:time,level,tags"])
doc.add_paragraph(
    "The log told a much harsher story than the 7.2 percent on screen. Over the monitored session "
    "the JVM ran 3,137 collections, 695 of them full GCs, and the pauses added up to 95.5 seconds. "
    "During the busiest ten second stretch the application was stopped for 68 percent of the time."
)
doc.add_paragraph(
    "It took me a while to work out why those two numbers disagree so badly, and the answer is "
    "that they measure different things. The serial collector pauses the application and then does "
    "its work on a single thread. This machine has 8 cores, and VisualVM reports GC activity as a "
    "share of total CPU capacity across all of them. One core fully busy out of eight is about "
    "12 percent of capacity, so 68 percent of wall time on one core lands near the 7 percent "
    "VisualVM displayed. The graph was not wrong, it just answers \"how much of my CPU is going "
    "into GC\" while the log answers \"how much of the time is my application frozen\". For a "
    "service that has to respond to requests, the second question is the one that matters."
)

# --------------------------------------------------------------- step 3
doc.add_heading("4. Step 3: Optimizing the JVM Settings", level=1)
doc.add_paragraph("Based on what VisualVM and the GC log showed, I changed the settings like this:")
table(
    ["Setting", "Before", "After", "Why"],
    [
        ["Initial heap", "-Xms64m", "-Xms512m",
         "Starting equal to the maximum stops the JVM from repeatedly growing the heap while it "
         "is already busy."],
        ["Maximum heap", "-Xmx128m", "-Xmx512m",
         "The real fix. There was never enough room for the short lived objects, so collections "
         "never stopped."],
        ["Collector", "-XX:+UseSerialGC", "-XX:+UseG1GC",
         "Serial GC uses one thread and stops the whole application. G1 works in parallel across "
         "cores and does much of its work concurrently."],
        ["Pause target", "not set", "-XX:MaxGCPauseMillis=100",
         "Tells G1 to aim for pauses under 100 ms, which it balances by resizing its regions."],
    ],
    widths=[1.1, 1.15, 1.15, 2.9],
)
doc.add_paragraph("The tuned command:")
code_block([
    "java -Xms512m -Xmx512m -XX:+UseG1GC -XX:MaxGCPauseMillis=100 \\",
    "     -cp \"target/classes:$(cat cp.txt)\" \\",
    "     com.example.itemservice.ItemserviceApplication",
])
figure("fig-tuned-overview.png",
       "Overview tab for the tuned run, confirming the new flags are in effect")

doc.add_paragraph(
    "I attached VisualVM again and ran exactly the same load. The Monitor tab looked like a "
    "different application."
)
figure("fig-tuned-monitor.png",
       "Monitor tab during the tuned run. Used heap now swings between roughly 100 MB and 400 MB "
       "inside the 512 MB space, with clear gaps between collections, and GC activity has dropped "
       "to 0.1 percent")
doc.add_paragraph(
    "Max heap now reads 536,870,912 B. Used heap rises and falls between about 100 MB and 400 MB "
    "with visible gaps instead of a solid band, which means the collector gets to wait between "
    "cycles rather than running constantly. GC activity fell from 7.2 percent to 0.1 percent, and "
    "CPU usage actually went up slightly, from 34 percent to 39 percent. That sounds wrong until "
    "you think about it: the CPU is now doing useful work answering requests instead of collecting "
    "garbage."
)

# --------------------------------------------------------------- results
doc.add_heading("5. Results", level=1)
doc.add_paragraph(
    "The two monitored sessions above ran the same load for the same length of time, so the GC "
    "logs from them can be compared directly."
)
table(
    ["Monitored session", "Baseline", "Tuned"],
    [
        ["Max heap (from VisualVM)", "134,217,728 B", "536,870,912 B"],
        ["GC activity (from VisualVM)", "7.2%", "0.1%"],
        ["CPU usage (from VisualVM)", "34.4%", "39.0%"],
        ["Collections", "3,137", "458"],
        ["Full GCs", "695", "0"],
        ["Total stop-the-world pause", "95.5 s", "2.3 s"],
        ["Pause share, busiest 10 s", "68%", "3%"],
    ],
    widths=[2.6, 1.9, 1.9],
)

doc.add_paragraph(
    "Screenshots show behaviour but not throughput, so I also ran the load test on its own with "
    "VisualVM detached, 1,200 requests with 16 concurrent threads at 25,000 items each. Same code, "
    "same machine. The only difference between the columns is the JVM flags."
)
table(
    ["Measurement", "Before tuning", "After tuning", "Change"],
    [
        ["Requests handled per second", "103.0", "173.5", "68% faster"],
        ["Average response time", "155.0 ms", "92.0 ms", "41% lower"],
        ["Median (p50) response time", "161.2 ms", "92.6 ms", "43% lower"],
        ["95th percentile response time", "287.7 ms", "115.3 ms", "60% lower"],
        ["Slowest single request", "454.3 ms", "173.1 ms", "62% lower"],
        ["Total time for 1,200 requests", "11.65 s", "6.92 s", "4.73 s saved"],
        ["Collections", "369", "32", "91% fewer"],
        ["Full GCs", "76", "0", "eliminated"],
        ["Total stop-the-world pause", "9,812 ms", "191 ms", "98% lower"],
    ],
    widths=[2.2, 1.4, 1.4, 1.3],
)
doc.add_paragraph(
    "The full GC count is the number I find most telling. A full GC stops every application thread "
    "while it collects the entire heap, and there were 76 of them before the change and none at "
    "all after. The 95th percentile dropping from 287.7 ms to 115.3 ms is the same story seen from "
    "the user's side, because those long pauses were exactly what made the slowest requests slow."
)
doc.add_paragraph(
    "One detail worth being careful about: the 9,812 ms of pause was measured across the 18.9 "
    "seconds that the GC log covers, which includes the warm-up requests as well as the measured "
    "ones, so it works out to 52 percent of that window rather than the 84 percent I first "
    "calculated by dividing it by the 11.65 second measured phase alone. I also had to separate "
    "G1's concurrent phases from its stop-the-world pauses when adding up the tuned figure, since "
    "concurrent work does not freeze the application and counting it would have overstated the "
    "pause total by about 45 ms."
)

doc.add_heading("A result I did not expect", level=2)
doc.add_paragraph(
    "Before the heavy test I had run a lighter one, 2,000 requests of 5,000 items each with 8 "
    "threads. Those numbers were much less exciting:"
)
table(
    ["Measurement", "Before tuning", "After tuning"],
    [
        ["Requests per second", "577.2", "571.3"],
        ["Average response time", "13.8 ms", "14.0 ms"],
        ["Collections", "160", "14"],
        ["Full GCs", "2", "0"],
        ["Total stop-the-world pause", "483 ms", "102 ms"],
    ],
    widths=[2.4, 1.7, 1.7],
)
doc.add_paragraph(
    "Throughput did not improve at all here. It actually came out 1 percent lower, which is small "
    "enough to just be noise between runs. GC work still dropped a lot, from 160 collections to "
    "14, but the application was not waiting on garbage collection in the first place at this "
    "level of load, so fixing GC gave nothing back in speed. I only saw the 68 percent gain once "
    "the load was heavy enough to actually exhaust the small heap."
)

# ------------------------------------------------------- lessons learned
doc.add_heading("6. Lessons Learned", level=1)
for text in [
    "Tuning only helps when the thing you are tuning is the actual bottleneck. My light load test "
    "proved that. GC activity fell sharply and the response times did not budge, because garbage "
    "collection was not what was holding it back at that point.",

    "Measure before changing anything. If I had jumped straight to changing flags I would have had "
    "no way of knowing whether it helped, or which flag was responsible.",

    "Check what a percentage is actually a percentage of. The 7.2 percent GC activity in VisualVM "
    "looked harmless next to a log showing the application frozen 68 percent of the time. Both "
    "were correct, because one is a share of eight cores and the other is a share of elapsed time.",

    "Graphs and logs answer different questions. VisualVM showed me where to look, since I could "
    "see the heap pinned at its limit and the thread count sitting flat. For numbers I could put "
    "in a table, the GC log was far better than reading values off a moving chart.",

    "Full GC count is a more useful warning sign than heap usage. A heap sitting near its limit is "
    "not automatically a problem, that can just be the collector being efficient. Repeated full "
    "GCs are, because each one freezes the whole application.",

    "Rising CPU can be a good sign. CPU went up from 34 to 39 percent after tuning, which looks "
    "like a regression until you notice throughput went up 68 percent at the same time. The CPU "
    "was finally being spent on requests instead of on garbage collection.",

    "A bigger heap is not automatically better. I went to 512 MB because that was the size that "
    "removed the full GCs. Setting it enormously high would have wasted memory the service does "
    "not need, and very large heaps can make individual collections take longer.",

    "The default settings on my machine would have hidden all of this. Java chose a 7.9 GB maximum "
    "heap and G1 on its own, and the endpoint never struggled. I had to deliberately restrict the "
    "heap to 128 MB to create a problem worth investigating, which is roughly the situation you "
    "would be in deploying to a small container with a memory limit.",
]:
    doc.add_paragraph(text, style="List Bullet")

# --------------------------------------------------------------- summary
doc.add_heading("7. Conclusion", level=1)
doc.add_paragraph(
    "I built a Spring Boot microservice with one REST endpoint, monitored it in VisualVM while it "
    "was under load, found that a 128 MB heap with the serial collector was leaving the application "
    "frozen for most of the test, and fixed it by raising the heap to 512 MB and switching to G1 "
    "with a 100 ms pause target. That took throughput from 103 to 173.5 requests per second and "
    "removed all 76 full GCs. The parts I will remember are that the same change made no "
    "difference under lighter load, and that the most alarming-looking number on the VisualVM "
    "dashboard was not the one that mattered."
)

# -------------------------------------------------------------- appendix
doc.add_heading("Appendix: Reproducing the Tests", level=1)
doc.add_paragraph("Baseline run:")
code_block([
    "java -Xms64m -Xmx128m -XX:+UseSerialGC \\",
    "     -Xlog:gc:file=baseline-gc.log:time,level,tags \\",
    "     -jar target/itemservice-0.0.1-SNAPSHOT.jar",
])
doc.add_paragraph("Tuned run:")
code_block([
    "java -Xms512m -Xmx512m -XX:+UseG1GC -XX:MaxGCPauseMillis=100 \\",
    "     -Xlog:gc:file=tuned-gc.log:time,level,tags \\",
    "     -jar target/itemservice-0.0.1-SNAPSHOT.jar",
])
doc.add_paragraph(
    "Load came from loadtest.py (1,200 requests, 16 concurrent, count=25000). Counting the "
    "collections and the stop-the-world pauses afterwards, note that matching on \"Pause\" is what "
    "keeps G1's concurrent phases out of the total:"
)
code_block([
    "grep -c 'Pause Young\\|Pause Full' baseline-gc.log",
    "grep -c 'Pause Full' baseline-gc.log",
    "grep -oP 'Pause (Young|Full).*?\\K\\d+\\.\\d+(?=ms)' baseline-gc.log \\",
    "    | awk '{s+=$1} END {print s \" ms\"}'",
])

doc.save(OUT)
print("wrote", OUT)
