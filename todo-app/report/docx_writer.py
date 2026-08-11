"""Writes the blocks out as a .docx."""
from __future__ import annotations

import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

PRINTABLE_IN = 6.5


def _shade(paragraph, fill: str) -> None:
    pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    pr.append(shd)


def _border(paragraph, colour: str = "C9CEDA") -> None:
    pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:space"), "6")
        el.set(qn("w:color"), colour)
        borders.append(el)
    pr.append(borders)


def write(blocks, out_path: str, fig_dir: str, meta) -> None:
    doc = Document()

    for section in doc.sections:
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.10

    # ---------------------------------------------------------------- heading
    day = doc.add_paragraph()
    day.alignment = WD_ALIGN_PARAGRAPH.CENTER
    day.paragraph_format.space_after = Pt(2)
    run = day.add_run(meta.DAY)
    run.bold = True
    run.font.size = Pt(12)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run(meta.TITLE)
    run.bold = True
    run.font.size = Pt(16)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(16)
    for text, more in [(meta.AUTHOR, True), (meta.COURSE, True),
                       (meta.DATE, False)]:
        r = sub.add_run(text)
        r.font.size = Pt(10.5)
        if more:
            r.add_break()

    fig_no = 0

    for block in blocks:
        kind = block[0]

        if kind == "h1":
            doc.add_heading(block[1], level=1)

        elif kind == "p":
            doc.add_paragraph(block[1])

        elif kind == "code":
            lines = block[1]
            for i, line in enumerate(lines):
                para = doc.add_paragraph()
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = (
                    Pt(0) if i < len(lines) - 1 else Pt(10))
                para.paragraph_format.left_indent = Inches(0.25)
                r = para.add_run(line if line else " ")
                r.font.name = "Consolas"
                r.font.size = Pt(9.5)
                _shade(para, "F2F2F2")

        elif kind == "fig":
            _, filename, caption, width = block
            width = min(width, PRINTABLE_IN)
            fig_no += 1
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(8)
            para.paragraph_format.space_after = Pt(2)
            para.add_run().add_picture(os.path.join(fig_dir, filename),
                                       width=Inches(width))
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.paragraph_format.space_after = Pt(14)
            r = cap.add_run(f"Figure {fig_no}: {caption}")
            r.italic = True
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

        elif kind == "table":
            rows, widths = block[1], block[2]
            table = doc.add_table(rows=1, cols=len(widths))
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            # Without a fixed layout Word ignores the column widths and
            # stretches the table across the page.
            table.autofit = False
            twips = [int(w * 1440) for w in widths]
            grid = table._tbl.find(qn("w:tblGrid"))
            for col, tw in zip(grid.findall(qn("w:gridCol")), twips):
                col.set(qn("w:w"), str(tw))

            header = table.rows[0].cells
            for i, text in enumerate(rows[0]):
                header[i].text = ""
                r = header[i].paragraphs[0].add_run(text)
                r.bold = True
                r.font.size = Pt(10)
            for row in rows[1:]:
                cells = table.add_row().cells
                for i, text in enumerate(row):
                    cells[i].text = ""
                    r = cells[i].paragraphs[0].add_run(str(text))
                    r.font.size = Pt(10)
            for row in table.rows:
                for i, tw in enumerate(twips):
                    row.cells[i].width = Inches(tw / 1440)
            doc.add_paragraph()

        elif kind == "bullets":
            for item in block[1]:
                doc.add_paragraph(item, style="List Bullet")

        elif kind == "note":
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(12)
            para.paragraph_format.left_indent = Inches(0.1)
            para.paragraph_format.right_indent = Inches(0.1)
            r = para.add_run(block[1])
            r.font.size = Pt(10)
            _shade(para, "F4F6FA")
            _border(para)

        elif kind == "break":
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

        else:
            raise ValueError(f"unknown block: {kind}")

    doc.save(out_path)
    return fig_no
